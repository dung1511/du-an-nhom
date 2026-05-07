from __future__ import annotations

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt

BASE_DIR = Path(__file__).resolve().parents[1]
DOC_PATH = BASE_DIR / 'docs' / 'nguoi2_booking_rooms_core_chi_tiet.docx'


def read_text(path: Path) -> str:
    return path.read_text(encoding='utf-8')


def find_line_number(text: str, needle: str) -> int:
    for idx, line in enumerate(text.splitlines(), start=1):
        if line.strip().startswith(needle):
            return idx
    return -1


def extract_top_level_block(text: str, starter: str) -> str:
    lines = text.splitlines()
    start = -1
    for i, line in enumerate(lines):
        if line.startswith(starter):
            start = i
            break
    if start == -1:
        return f'# Khong tim thay: {starter}'

    end = len(lines)
    for j in range(start + 1, len(lines)):
        if lines[j].startswith('def ') or lines[j].startswith('class '):
            end = j
            break

    return '\n'.join(lines[start:end]).rstrip() + '\n'


def add_code_section(doc: Document, file_rel: str, symbol: str, explanation: str, extractor='auto') -> None:
    file_path = BASE_DIR / file_rel
    text = read_text(file_path)

    if extractor == 'whole_file':
        code = text
        line_no = 1
        shown_symbol = 'Toan bo file'
    else:
        starter = symbol
        code = extract_top_level_block(text, starter)
        line_no = find_line_number(text, starter)
        shown_symbol = symbol

    doc.add_heading(f'{file_rel} - {shown_symbol} (dong {line_no if line_no > 0 else "N/A"})', level=3)
    doc.add_paragraph('Giai thich chi tiet:')
    doc.add_paragraph(explanation)
    doc.add_paragraph('Code:')
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)


def generate() -> None:
    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    doc.add_heading('Nguoi 2 - Booking + Rooms (Core chinh) - Tai lieu day du code + giai thich', level=1)
    doc.add_paragraph(
        'Tai lieu tong hop toan bo doan code lien quan trong rooms/ va templates/rooms/ '
        'cho cac chuc nang: hien thi danh sach phong, chi tiet phong, booking, upload anh phong, '
        'trang thai phong con/het, lich dat phong, admin duyet don va CRUD lich su dat phong.'
    )

    doc.add_heading('A. Rooms Core (models + manager)', level=2)
    add_code_section(
        doc,
        'rooms/models.py',
        'def room_cover_upload_path(instance, filename):',
        'Ham tao ten file anh dai dien phong bang UUID, tranh trung ten va tranh loi overwrite file.'
    )
    add_code_section(
        doc,
        'rooms/models.py',
        'def room_gallery_upload_path(instance, filename):',
        'Ham tao duong dan upload gallery anh phong. Dung cho chuc nang upload nhieu anh phong.'
    )
    add_code_section(
        doc,
        'rooms/models.py',
        'class RoomManager(models.Manager):',
        'Manager chua cac ham tim phong trong, tim phong phu hop va goi y to hop phong cho nhom dong.'
    )
    add_code_section(
        doc,
        'rooms/models.py',
        'class Room(models.Model):',
        'Model phong va cac ham cot loi: can_accommodate, is_available, availability_status. '
        'Day la trung tam cua logic trang thai phong con/het va suc chua.'
    )
    add_code_section(
        doc,
        'rooms/models.py',
        'class RoomImage(models.Model):',
        'Model luu gallery anh phong, lien ket 1-n voi Room qua khoa ngoai room.'
    )
    add_code_section(
        doc,
        'rooms/models.py',
        'class Reservation(models.Model):',
        'Model dat phong: thong tin khach, check-in/check-out, thanh toan, coc, hu hong, lich su don. '
        'Day la du lieu nen cho booking va lich su dat phong.'
    )

    doc.add_heading('B. API + Web views cho Booking/Rooms', level=2)
    for file_rel, symbol, explanation in [
        ('rooms/views.py', 'class RoomListAPIView(generics.ListAPIView):', 'API lay danh sach phong, co filter keyword/category/size/price/rating va phan trang.'),
        ('rooms/views.py', 'class RoomDetailAPIView(generics.RetrieveAPIView):', 'API chi tiet 1 phong theo id.'),
        ('rooms/views.py', 'class RoomSearchAPIView(APIView):', 'API tim phong theo ngay + so khach, tra ket qua phong phu hop va goi y combo.'),
        ('rooms/views.py', 'def room_list(request):', 'View HTML hien thi danh sach phong theo danh muc.'),
        ('rooms/views.py', 'def room_detail(request, room_id):', 'View HTML chi tiet phong + feedback form.'),
        ('rooms/views.py', 'def room_search(request):', 'View HTML tim kha dung phong theo ngay, suc chua, va goi y phong/combos.'),
        ('rooms/views.py', 'def room_booking(request):', 'View HTML booking chinh: tinh tien, coupon, dich vu, tao reservation.'),
        ('rooms/views.py', 'def booking_confirmation(request, reservation_id):', 'Trang xac nhan booking, dung context hoa don chi tiet.'),
        ('rooms/views.py', 'def my_bookings(request):', 'Trang lich su booking cua user dang nhap.'),
        ('rooms/views.py', 'def cancel_reservation(request, reservation_id):', 'Huy booking online neu con trong thoi han cho phep; cap nhat trang thai lich su.'),
        ('rooms/views.py', 'class ReservationListCreateAPIView(generics.ListCreateAPIView):', 'API tao booking cho user da dang nhap va gui email sau khi commit.'),
        ('rooms/views.py', 'class ReservationListAPIView(generics.ListAPIView):', 'API list reservation cua user va POST tao reservation qua serializer.'),
        ('rooms/views.py', 'class ReservationDetailAPIView(generics.RetrieveAPIView):', 'API xem chi tiet reservation theo quyen user/admin.'),
        ('rooms/views.py', 'class ReservationCheckedOutListAPIView(generics.ListAPIView):', 'API xem lich su don da checkout.'),
    ]:
        add_code_section(doc, file_rel, symbol, explanation)

    doc.add_heading('C. Upload anh phong + bao mat upload', level=2)
    add_code_section(
        doc,
        'rooms/views.py',
        'class RoomImageUploadAPIView(APIView):',
        'API admin upload nhieu anh phong: validate room_id, gioi han so anh, transaction rollback neu loi.'
    )
    add_code_section(
        doc,
        'rooms/upload_security.py',
        'def validate_image_file(uploaded_file, max_size=MAX_UPLOAD_SIZE):',
        'Kiem tra kich thuoc, mime type, extension va verify anh bang PIL de ngan file doc hai.'
    )
    add_code_section(
        doc,
        'rooms/upload_security.py',
        'def build_safe_filename(extension):',
        'Sinh ten file an toan bang UUID + extension da chuan hoa.'
    )
    add_code_section(
        doc,
        'rooms/views.py',
        'def admin_room_image_upload_page(request):',
        'Trang web admin de chon phong va upload gallery anh.'
    )

    doc.add_heading('D. Trang thai phong con/het + lich dat phong', level=2)
    add_code_section(
        doc,
        'rooms/views.py',
        'def check_room_availability_api(request):',
        'API nhanh kiem tra phong co bi dat trung ngay hay khong (true/false).'
    )
    add_code_section(
        doc,
        'rooms/views.py',
        'def room_list_filtered(request):',
        'Tim phong theo check-in/check-out + so khach va render danh sach phu hop + combo.'
    )
    add_code_section(
        doc,
        'rooms/views.py',
        'def room_combo_detail(request):',
        'Trang chi tiet phuong an ghep nhieu phong cho nhom dong.'
    )

    doc.add_heading('E. Admin duyet don + them/sua/xoa lich su dat phong', level=2)
    add_code_section(
        doc,
        'rooms/admin.py',
        'class ReadOnlyForStaffAdminMixin:',
        'Phan quyen CRUD tren Django Admin: staff xem, superuser moi duoc them/sua/xoa.'
    )
    add_code_section(
        doc,
        'rooms/admin.py',
        'class ReservationAdmin(ReadOnlyForStaffAdminMixin, admin.ModelAdmin):',
        'Man hinh admin trung tam quan ly lich su booking: list_display, filter, search, readonly_fields, fields, actions.'
    )
    add_code_section(
        doc,
        'rooms/admin.py',
        'def confirm_deposit_action(self, request, queryset):',
        'Action duyet bien lai coc hang loat tren admin.'
    )
    add_code_section(
        doc,
        'rooms/views.py',
        'class AdminConfirmDepositAPIView(APIView):',
        'API admin xac nhan bien lai coc cho 1 reservation.'
    )
    add_code_section(
        doc,
        'rooms/views.py',
        'class AdminReservationListAPIView(generics.ListAPIView):',
        'API admin xem toan bo lich su reservation (phan trang).'
    )
    add_code_section(
        doc,
        'rooms/views.py',
        'class AdminCheckedOutReservationListAPIView(generics.ListAPIView):',
        'API admin xem lich su da checkout + filter room/user.'
    )

    doc.add_heading('F. Serializers cot loi (validation + output)', level=2)
    for file_rel, symbol, explanation in [
        ('rooms/serializers.py', 'class ReservationCreateSerializer(serializers.ModelSerializer):', 'Validation booking: ngay, suc chua, phong trong, coupon, user gan booking. Tinh toan tien ban dau.'),
        ('rooms/serializers.py', 'class ReservationDetailSerializer(serializers.ModelSerializer):', 'Serializer tra thong tin reservation chi tiet cho lich su/chi tiet don.'),
        ('rooms/serializers.py', 'class RoomSerializer(serializers.ModelSerializer):', 'Serializer phong co them availability_status va is_available_now de hien thi con/het.'),
        ('rooms/serializers.py', 'class RoomSearchSerializer(serializers.Serializer):', 'Validation dau vao cho tim phong theo lich.'),
        ('rooms/serializers.py', 'class UploadDepositReceiptSerializer(serializers.ModelSerializer):', 'Serializer nhan file bien lai coc tu khach.'),
    ]:
        add_code_section(doc, file_rel, symbol, explanation)

    doc.add_heading('G. URL map tong hop luong Rooms/Booking', level=2)
    add_code_section(
        doc,
        'rooms/urls.py',
        'urlpatterns = [',
        'Danh sach URL map cho API va web view rooms/booking/admin/frontdesk.',
        extractor='whole_file'
    )

    doc.add_heading('H. Templates Rooms (giao dien chuc nang)', level=2)
    for tpl, summary in [
        ('templates/rooms/rooms.html', 'Giao dien danh sach phong.'),
        ('templates/rooms/roomdetail.html', 'Giao dien chi tiet phong va hanh dong booking.'),
        ('templates/rooms/roomsearch.html', 'Giao dien ket qua tim phong theo lich.'),
        ('templates/rooms/roomsfilter.html', 'Giao dien danh sach phong loc theo ngay/so khach.'),
        ('templates/rooms/roombooking.html', 'Form booking va tong hop thanh toan.'),
        ('templates/rooms/bookingconfirmation.html', 'Trang xac nhan booking va hoa don tom tat.'),
        ('templates/rooms/mybookings.html', 'Trang lich su booking cua nguoi dung.'),
        ('templates/rooms/admin_room_image_upload.html', 'Trang admin upload anh gallery phong.'),
        ('templates/rooms/upload_deposit.html', 'Trang khach upload bien lai coc qua QR.'),
        ('templates/rooms/frontdesk_dashboard.html', 'Trang le tan xu ly check-in/check-out + tra cuu.'),
        ('templates/rooms/frontdesk_print_slip.html', 'Mau phieu in cho le tan va hoa don cap nhat.'),
    ]:
        path = BASE_DIR / tpl
        if path.exists():
            add_code_section(doc, tpl, '', summary, extractor='whole_file')

    doc.add_heading('I. Ket luan hoc tap', level=2)
    doc.add_paragraph(
        'Thu tu hoc de nhanh hieu he thong:\n'
        '1) models Room/Reservation -> 2) serializer create/search -> 3) views room_list/room_detail/room_booking ->\n'
        '4) templates rooms/roombooking/mybookings -> 5) admin ReservationAdmin + API admin.'
    )

    doc.save(DOC_PATH)
    print(f'WROTE:{DOC_PATH}')


if __name__ == '__main__':
    generate()
