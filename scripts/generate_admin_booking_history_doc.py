from docx import Document
from docx.shared import Pt
import os

DOC_PATH = os.path.join('docs', 'admin_booking_history_crud_explanation.docx')

sections = [
    (
        'Muc tieu tai lieu',
        'Tai lieu nay chi tap trung vao phan ADMIN cho lich su dat phong: Them, Sua, Xoa va xem lich su.\n'
        'Noi dung giai thich theo vi tri code that trong du an de ban hoc truc tiep vao dung cho.'
    ),
    (
        '1) Code goc cua admin quan ly lich su dat phong nam o dau',
        'File chinh: rooms/admin.py\n'
        '- rooms/admin.py:15 class ReadOnlyForStaffAdminMixin\n'
        '- rooms/admin.py:120 class ReservationAdmin\n'
        '- rooms/admin.py:236 def confirm_deposit_action\n\n'
        'File route API lien quan (ho tro dashboard/lich su): rooms/urls.py\n'
        '- rooms/urls.py:26 api/admin/reservations/\n'
        '- rooms/urls.py:28 api/admin/checked-out-reservations/\n'
        '- rooms/urls.py:27 api/admin/reservations/<id>/confirm-deposit/\n\n'
        'File view API admin lien quan: rooms/views.py\n'
        '- rooms/views.py:966 class AdminReservationListAPIView\n'
        '- rooms/views.py:975 class AdminCheckedOutReservationListAPIView\n'
        '- rooms/views.py:434 class AdminConfirmDepositAPIView'
    ),
    (
        '2) Them/Sua/Xoa o Admin duoc quyen nhu the nao',
        'Quyen duoc dat trong mixin rooms/admin.py:15 (ReadOnlyForStaffAdminMixin):\n\n'
        'A. has_view_permission (admin.py:18):\n'
        '- Staff hoac superuser deu duoc xem.\n\n'
        'B. has_add_permission (admin.py:30):\n'
        '- CHI superuser moi duoc Them reservation trong trang admin.\n\n'
        'C. has_change_permission (admin.py:38):\n'
        '- CHI superuser moi duoc Sua reservation.\n\n'
        'D. has_delete_permission (admin.py:46):\n'
        '- CHI superuser moi duoc Xoa reservation.\n\n'
        'Ket luan nhanh:\n'
        '- Staff: xem duoc, nhung khong them/sua/xoa.\n'
        '- Superuser: day du CRUD (create/read/update/delete).'
    ),
    (
        '3) ReservationAdmin quan ly lich su dat phong chi tiet',
        'Vi tri: rooms/admin.py:120 class ReservationAdmin\n\n'
        'Cac thanh phan quan trong:\n'
        '- list_display (admin.py:124): hien thi cot user, room, check_in_date, check_out_date, is_checked_out, total, deposit_status_label, created_at, checkout_at.\n'
        '- list_filter (admin.py:137): loc theo room, ngay, payment_method, checkout_at, trang thai.\n'
        '- search_fields (admin.py:145): tim theo ten khach, email, ten phong.\n'
        '- readonly_fields (admin.py:159): khoi field quan trong tranh sua tay sai quy trinh (vd deposit_confirmed_at, deposit_confirmed_by, checkout_at...).\n'
        '- fields (admin.py:171): sap xep thu tu field trong form admin de admin thao tac de doc/de kiem soat.\n\n'
        'Y nghia hoc tap:\n'
        '- Day la cho trung tam de quan ly lich su dat phong tren giao dien admin cua Django (khong can viet giao dien rieng).'
    ),
    (
        '4) Lich su dat phong duoc luu boi truong nao trong model',
        'Vi tri: rooms/models.py\n'
        '- rooms/models.py:807 class Reservation\n'
        '- rooms/models.py:878 created_at (thoi gian tao don)\n'
        '- rooms/models.py:957 is_checked_in\n'
        '- rooms/models.py:960 checked_in_at\n'
        '- rooms/models.py:972 is_checked_out\n'
        '- rooms/models.py:1117 checkout_at\n'
        '- rooms/models.py:1133 is_canceled\n'
        '- rooms/models.py:1136 canceled_at\n'
        '- rooms/models.py:1223 deposit_confirmed\n'
        '- rooms/models.py:1226 deposit_confirmed_at\n'
        '- rooms/models.py:1232 deposit_confirmed_by\n\n'
        'Nghia la “lich su” khong chi la 1 bang log rieng, ma la bo trang thai + moc thoi gian tren Reservation.'
    ),
    (
        '5) Admin duyet coc (mot phan cua xu ly don trong admin)',
        'Co 2 cach duyet trong code:\n\n'
        'A. Action trong Django Admin:\n'
        '- rooms/admin.py:236 confirm_deposit_action\n'
        '- Admin chon nhieu booking -> action xac nhan coc hang loat\n'
        '- He thong cap nhat: deposit_confirmed, deposit_confirmed_at, deposit_confirmed_by, payment_status\n\n'
        'B. API duyet tu he thong ngoai/admin panel rieng:\n'
        '- rooms/views.py:434 AdminConfirmDepositAPIView\n'
        '- URL: rooms/urls.py:27\n\n'
        'Diem hoc:\n'
        '- Du cung 1 nghiep vu, co the cung cap ca giao dien Django Admin va API rieng.'
    ),
    (
        '6) API “lich su dat phong” ben admin hoc o dau',
        'A. Danh sach toan bo reservation (admin):\n'
        '- rooms/views.py:966 AdminReservationListAPIView\n'
        '- rooms/urls.py:26 api/admin/reservations/\n\n'
        'B. Danh sach reservation da tra phong (lich su hoan tat):\n'
        '- rooms/views.py:975 AdminCheckedOutReservationListAPIView\n'
        '- rooms/urls.py:28 api/admin/checked-out-reservations/\n\n'
        'Luu y quan trong:\n'
        '- Cac API tren hien tai la LIST (xem/loc) chu KHONG phai endpoint CRUD day du.\n'
        '- CRUD chinh dang dua vao Django Admin + quyen superuser trong admin.py.'
    ),
    (
        '7) Neu ban can “day du CRUD qua API” (goi y nang cao)',
        'Hien tai code admin API chua co endpoint tao/sua/xoa reservation cho admin.\n'
        'Neu muon hoc nang cao, ban co the them:\n'
        '- POST /api/admin/reservations/ (tao)\n'
        '- PATCH /api/admin/reservations/<id>/ (sua)\n'
        '- DELETE /api/admin/reservations/<id>/ (xoa)\n'
        'kem permission IsStaffOrAdmin + audit log nguoi sua/xoa.'
    ),
    (
        '8) Tom tat nhanh de nho',
        'Noi Them/Sua/Xoa lich su dat phong o ben admin “dang nam o day”:\n'
        '- rooms/admin.py:15, 30, 38, 46 (permission CRUD)\n'
        '- rooms/admin.py:120 (ReservationAdmin giao dien quan ly)\n'
        '- rooms/admin.py:236 (action duyet coc hang loat)\n\n'
        'Noi xem lich su qua API admin:\n'
        '- rooms/views.py:966, 975\n'
        '- rooms/urls.py:26, 28'
    ),
]


def make_doc(path: str) -> None:
    os.makedirs(os.path.dirname(path), exist_ok=True)

    doc = Document()
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)

    doc.add_heading('Giai thich code Admin: Them/Sua/Xoa lich su dat phong', level=1)

    for title, body in sections:
        doc.add_heading(title, level=2)
        for line in body.split('\n'):
            doc.add_paragraph(line)

    doc.save(path)
    print(f'WROTE:{path}')


if __name__ == '__main__':
    make_doc(DOC_PATH)
